import os
import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import sys

def convert_paragraph(element, doc):
    p = doc.add_paragraph()
    p.add_run(element.get_text()).font.size = Pt(12)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def convert_heading(element, doc, level):
    h = doc.add_heading(level=level)
    h.add_run(element.get_text())

def convert_table(element, doc):
    rows = element.find_all('tr')
    if rows:
        cols = rows[0].find_all(['td', 'th'])
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'  # Add table style
        hdr_cells = table.rows[0].cells
        for i, cell in enumerate(cols):
            hdr_cells[i].text = cell.get_text()
        for row in rows[1:]:
            cells = row.find_all(['td', 'th'])
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                row_cells[i].text = cell.get_text()

def markdown_to_word(markdown_text):
    # Convert markdown text to HTML with tables extension enabled
    html = markdown.markdown(markdown_text, extensions=['tables'])
    
    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create a new Document
    doc = Document()
    
    # Add HTML content to the document
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table']):
        if element.name == 'p':
            convert_paragraph(element, doc)
        elif element.name.startswith('h') and len(element.name) == 2:
            level = int(element.name[1])
            convert_heading(element, doc, level)
        elif element.name == 'table':
            convert_table(element, doc)
    
    return doc

def main(input_file, output_file):
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Convert markdown to word document
    doc = markdown_to_word(markdown_text)
    
    # Save the document
    doc.save(output_file)
    print(f"Markdown file '{input_file}' has been converted to Word document '{output_file}'")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python markdown_to_word.py <input_markdown_file> <output_word_file>")
        input_file = os.path.join('data', 'table1.md')
        output_file = os.path.join('output', 'table1.docx')
        main(input_file, output_file)
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        main(input_file, output_file)

